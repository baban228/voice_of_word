from docx import Document
from docx.shared import Pt, RGBColor
import os
from path_to_word import Path

class Write_in_word:
    def __init__(self, filename=""):
        self.filename = filename
        self.doc = Document(filename)

    def write_paragraph(self, text, font_name='Times New Roman', font_size=14, heading_style='Heading 1'):
        heading = self.doc.add_paragraph()
        # Добавляем текст с настройкой стилей
        run = heading.add_run(text)

        # Настройка шрифта
        run.font.name = font_name
        run.font.size = Pt(font_size)

        # Устанавливаем стиль заголовка
        heading.style = heading_style

    def write_title(self, text, font_name='Times New Roman', font_size=24, font_color=RGBColor(0, 0, 0), heading_style='Heading 1'):
        # Создаем параграф для заголовка
        heading = self.doc.add_paragraph()

        # Добавляем текст с настройкой стилей
        run = heading.add_run(text)

        # Настройка шрифта
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.color.rgb = font_color

        # Устанавливаем стиль заголовка
        heading.style = heading_style

    def save(self):
        path = Path()
        self.doc.save(path.get_path_file(path))

    def open_docx_file(self, file_path):
        if not file_path:
            raise ValueError("Путь к файлу не может быть пустым")

        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Файл не найден: {file_path}")

        if not file_path.endswith('.docx'):
            raise ValueError(f"Файл должен быть документом Word (.docx): {file_path}")

        try:
            doc = Document(file_path)
            print("Файл успешно открыт")
            # Здесь вы можете работать с документом
        except Exception as e:
            print(f"Ошибка при открытии файла: {e}")