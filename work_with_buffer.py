import docx

class Work_with_buffer:
    def __init__(self, txt_file_path, docx_file_path):
        self.txt_file_path = txt_file_path
        self.docx_file_path = docx_file_path

    def read_txt_file(self):
        """Читает текст из .txt файла"""
        with open(self.txt_file_path, 'r', encoding='utf-8') as file:
            return file.read()

    def read_docx_file(self):
        """Читает текст из .docx файла"""
        doc = docx.Document(self.docx_file_path)
        doc_text = ''
        for para in doc.paragraphs:
            doc_text += para.text + '\n'
        return doc_text

    def append_text_to_docx(self, new_text):
        """Добавляет текст в конец .docx файла"""
        doc = docx.Document(self.docx_file_path)
        doc.add_paragraph(new_text)
        doc.save(self.docx_file_path)

    def compare_and_update(self):
        """Сравнивает файлы и добавляет текст, если его нет в .docx"""
        txt_text = self.read_txt_file()
        docx_text = self.read_docx_file()

        if txt_text not in docx_text:
            self.append_text_to_docx(txt_text)
            print("Текст из txt был добавлен в docx.")
        else:
            print("Текст из txt уже присутствует в docx.")
